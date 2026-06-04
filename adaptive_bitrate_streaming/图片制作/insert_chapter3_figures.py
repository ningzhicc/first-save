from pathlib import Path
from shutil import copy2

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.shared import Inches, Pt


BASE_DIR = Path("/data3/wangxh/NetLLM-master/adaptive_bitrate_streaming")
DOC_PATH = BASE_DIR / "第三章.docx"
BACKUP_PATH = BASE_DIR / "第三章_插图前备份.docx"
OUTPUT_PATH = BASE_DIR / "第三章_插图版.docx"
FIG_DIR = BASE_DIR / "论文书写" / "论文图片"

FIGURES = [
    {
        "section": "3.1",
        "next_section": "3.2",
        "image": FIG_DIR / "fig3_1_llm_abr_policy_overview_cn.png",
        "caption": "图 3.1 基于大语言模型适配的 ABR 策略模型总体结构",
        "width": 6.1,
    },
    {
        "section": "3.2",
        "next_section": "3.3",
        "image": FIG_DIR / "fig3_2_semantic_reprogramming_encoder_cn.png",
        "caption": "图 3.2 语义重编程状态编码示意图",
        "width": 6.0,
    },
    {
        "section": "3.4",
        "next_section": "3.5",
        "image": FIG_DIR / "fig3_3_context_alignment_multiscale_history_cn.png",
        "caption": "图 3.3 上下文预对齐与轻量化多尺度历史建模结构图",
        "width": 6.1,
    },
]


def paragraph_after(paragraph):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    return paragraph.__class__(new_p, paragraph._parent)


def insert_after(paragraph, text=None):
    new_para = paragraph_after(paragraph)
    if text is not None:
        new_para.add_run(text)
    return new_para


def find_section_range(doc, section_prefix, next_section_prefix):
    start = None
    end = None
    for idx, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if start is None and text.startswith(section_prefix):
            start = idx
            continue
        if start is not None and text.startswith(next_section_prefix):
            end = idx
            break
    if start is None or end is None:
        raise ValueError(f"Cannot find section range: {section_prefix} to {next_section_prefix}")
    return start, end


def apply_caption_style(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in paragraph.runs:
        run.font.name = "宋体"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        run.font.size = Pt(10.5)


def qn(tag):
    prefix, tagroot = tag.split(":")
    nsmap = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    return f"{{{nsmap[prefix]}}}{tagroot}"


def insert_figure_after_section(doc, item):
    if any(item["caption"] in para.text for para in doc.paragraphs):
        return False

    _, next_idx = find_section_range(doc, item["section"], item["next_section"])
    anchor = doc.paragraphs[next_idx - 1]

    blank_after_caption = insert_after(anchor)
    caption_para = insert_after(anchor, item["caption"])
    image_para = insert_after(anchor)
    blank_before_image = insert_after(anchor)

    image_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    image_para.add_run().add_picture(str(item["image"]), width=Inches(item["width"]))
    apply_caption_style(caption_para)

    # Keep the surrounding spacing modest so the figure reads as part of the section.
    for para in (blank_before_image, image_para, caption_para, blank_after_caption):
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after = Pt(6)

    return True


def main():
    for item in FIGURES:
        if not item["image"].exists():
            raise FileNotFoundError(item["image"])

    if not BACKUP_PATH.exists():
        copy2(DOC_PATH, BACKUP_PATH)

    doc = Document(DOC_PATH)
    for item in reversed(FIGURES):
        insert_figure_after_section(doc, item)

    doc.save(OUTPUT_PATH)
    doc.save(DOC_PATH)

    check_doc = Document(DOC_PATH)
    print(f"saved={DOC_PATH}")
    print(f"copy={OUTPUT_PATH}")
    print(f"backup={BACKUP_PATH}")
    print(f"inline_shapes={len(check_doc.inline_shapes)}")
    for para in check_doc.paragraphs:
        text = para.text.strip()
        if text.startswith("图 3."):
            print(text)


if __name__ == "__main__":
    main()
